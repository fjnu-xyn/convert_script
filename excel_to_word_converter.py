"""
Excel到Word转换脚本
功能：读取Excel文件，将内容按指定格式转换为Word文档
"""

import re
import os
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from logger import get_logger

logger = get_logger("excel_to_word_converter")

# 尝试导入验证模块，如果失败则忽略（兼容单独运行）
try:
    from verify_word import verify_consistency
except ImportError:
    verify_consistency = None


def set_font(run, font_name='宋体', font_size=10.5, bold=False):
    """设置字体格式"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = False  # 强制不倾斜
    # 设置中文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_styled_heading(doc, text, level):
    """添加带格式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        # 标题通常加粗，字号根据层级可能不同，这里统一按用户要求"黑色宋体"
        # Word默认标题字号较大，用户未指定字号，只说"黑色宋体"。
        # 为了美观，保持默认字号或稍微调整？
        # 用户之前的脚本里正文是10.5pt (五号)。
        # 标题通常比正文大。
        # 标题3: 16pt, 标题4: 14pt, 标题5: 12pt... 
        # 但用户要求"统一成黑色宋体"。
        # 让我们只强制字体和颜色，保留大小或设为合适的大小。
        # 既然用户强调"统一"，可能指字体类型。
        # 让我们显式设置一下颜色和字体名。
        set_font(run, font_name='宋体', font_size=get_font_size_for_level(level), bold=True)
    return heading

def get_font_size_for_level(level):
    """根据标题级别返回字号"""
    if level == 3: return 16  # 三号
    if level == 4: return 14  # 四号
    if level == 5: return 12  # 小四
    if level == 6: return 10.5 # 五号
    return 10.5


def split_subprocess_description(text):
    """
    拆分子过程描述字段
    
    仅按"输入-"、"查询-"、"呈现-"、"校验-"、"输出-"等前缀进行拆分
    忽略其他分隔符（逗号、句号等），保持每个前缀块的完整性
    
    Args:
        text: 子过程描述文本
        
    Returns:
        list: 拆分后的行列表
    """
    if not text or pd.isna(text):
        return []
    
    # 转换为字符串并清理
    text = str(text).strip()
    
    if not text:
        return []
    
    # 定义前缀词
    prefixes = ['输入-', '查询-', '呈现-', '校验-', '输出-']
    
    result = []
    current_segment = []
    
    # 按字符遍历，当遇到前缀时开始新段
    i = 0
    while i < len(text):
        # 检查当前位置是否是前缀开头
        is_prefix = False
        for prefix in prefixes:
            if text[i:i+len(prefix)] == prefix:
                # 保存前一个段落
                if current_segment:
                    result.append(''.join(current_segment).strip())
                    current_segment = []
                # 开始新段
                current_segment.append(prefix)
                i += len(prefix)
                is_prefix = True
                break
        
        if not is_prefix:
            current_segment.append(text[i])
            i += 1
    
    # 保存最后一个段落
    if current_segment:
        result.append(''.join(current_segment).strip())
    
    # 清理结果（移除空项和末尾分号）
    result = [seg.rstrip('；;').strip() for seg in result if seg]
    
    # 如果没有拆分出任何前缀段落（即不包含输入/查询等关键字），则返回原文本
    if not result and text:
        return [text]
        
    return result


def read_excel_robust(excel_path):
    """
    健壮地读取Excel文件，自动查找正确的Sheet和表头
    """
    try:
        xl = pd.ExcelFile(excel_path)
    except Exception as e:
        print(f"无法打开Excel文件: {e}")
        logger.error(f"无法打开Excel文件: {e}")
        return None

    # 1. 查找包含数据的Sheet
    target_sheet = None
    for sheet in xl.sheet_names:
        if '拆分表' in sheet or '功能点' in sheet:
            target_sheet = sheet
            break
    
    if not target_sheet:
        # 如果没找到特定名称的sheet，尝试使用最大的sheet（通常数据最多）
        # 或者默认使用第一个
        target_sheet = xl.sheet_names[0]
        print(f"未找到名称包含'拆分表'的Sheet，默认使用: {target_sheet}")
        logger.warning(f"未找到名称包含'拆分表'的Sheet，默认使用: {target_sheet}")
    else:
        print(f"使用Sheet: {target_sheet}")
        logger.info(f"使用Sheet: {target_sheet}")

    # 2. 查找表头行 - 处理多行表头
    # 读取前10行来分析
    df_preview = pd.read_excel(excel_path, sheet_name=target_sheet, header=None, nrows=10)
    
    # 尝试找到包含完整列信息的行
    # 策略：同时包含"客户需求"和"一级模块"的行，或包含"功能过程"和"子过程描述"的行
    header_candidates = []
    for idx, row in df_preview.iterrows():
        row_str = row.astype(str).values
        row_text = ' '.join(row_str)
        score = 0
        if '客户需求' in row_text: score += 1
        if '一级模块' in row_text: score += 1
        if '功能过程' in row_text or '功能名称' in row_text: score += 1
        if '子过程描述' in row_text or '功能描述' in row_text: score += 1
        if score >= 2:
            header_candidates.append((idx, score, row))
    
    if not header_candidates:
        print("未找到标准表头行，尝试使用多行表头策略")
        logger.info("未找到标准表头行，尝试使用多行表头策略")
        # 读取前3行作为多级表头
        df = pd.read_excel(excel_path, sheet_name=target_sheet, header=[0, 1, 2])
        # 合并多级列名
        df.columns = [' '.join([str(c).strip() for c in col if 'Unnamed' not in str(c)]).strip() 
                      for col in df.columns]
    else:
        # 使用得分最高的行作为表头
        header_candidates.sort(key=lambda x: x[1], reverse=True)
        header_row_idx = header_candidates[0][0]
        print(f"定位到表头在第 {header_row_idx} 行 (得分: {header_candidates[0][1]})")
        logger.info(f"定位到表头在第 {header_row_idx} 行 (得分: {header_candidates[0][1]})")
        df = pd.read_excel(excel_path, sheet_name=target_sheet, header=header_row_idx)
    
    return df


def excel_to_word(excel_path, word_path=None, perform_verify=True, open_output=True):
    """
    将Excel文件转换为Word文档
    :param open_output: 转换完成后是否自动打开文件（服务器模式下应设为False）
    """
    print(f"正在处理: {excel_path.name}")
    logger.info(f"正在处理: {excel_path.name}")
    
    # 读取Excel文件
    df = read_excel_robust(excel_path)
    if df is None:
        return

    # 自动识别列索引 - 增强模糊匹配能力
    # 我们需要找到: 客户需求, 一级模块, 二级模块, 三级模块, 功能过程, 子过程描述
    col_map = {}
    required_cols = {
        'CustomerReq': ['客户需求'],
        'Level1': ['一级模块'],
        'Level2': ['二级模块'],
        'Level3': ['三级模块'],
        'Process': ['功能过程'],
        'Description': ['子过程描述']
    }
    
    # 先尝试从列名精确匹配(优先级高)
    for col_idx, col_name in enumerate(df.columns):
        col_str = str(col_name).strip()
        # 精确匹配优先
        for key, keywords in required_cols.items():
            if key not in col_map:
                # 使用完全匹配避免误匹配
                if col_str in keywords:
                    col_map[key] = col_name
                    break
    
    # 检查是否找到所有列
    missing_cols = [k for k in required_cols if k not in col_map]
    
    # 特殊处理：检查第一行数据是否包含"一级模块"等信息（转置表头）
    if missing_cols and 'Level1' in missing_cols and len(df) > 0:
        first_row = df.iloc[0]
        
        # 检查第一行是否包含模块信息
        for idx, val in enumerate(first_row.values):
            val_str = str(val).strip()
            if '一级模块' in val_str and 'Level1' not in col_map:
                col_map['Level1'] = df.columns[idx]
            elif '二级模块' in val_str and 'Level2' not in col_map:
                col_map['Level2'] = df.columns[idx]
            elif '三级模块' in val_str and 'Level3' not in col_map:
                col_map['Level3'] = df.columns[idx]
        
        # 删除包含"模块"的元数据行
        df = df[~df.apply(lambda row: any('级模块' in str(val) for val in row.values), axis=1)].reset_index(drop=True)
        
    # 重新检查缺失列
    missing_cols = [k for k in required_cols if k not in col_map]
    
    # 如果没找到，尝试按固定索引回退 (针对 cosmic 表格结构)
    if missing_cols:
        print(f"警告: 未能通过列名自动识别所有列: {missing_cols}，尝试使用固定列索引策略...")
        logger.warning(f"未能通过列名自动识别所有列: {missing_cols}，尝试使用固定列索引策略...")
        
        # 检查列数是否足够
        if len(df.columns) >= 8:
            # 假设结构: [0]CustomerReq, [1]L1, [2]L2, [3]L3, ..., [6]Process, [7]Desc
            # 强制指定
            if 'CustomerReq' not in col_map: col_map['CustomerReq'] = df.columns[0]
            if 'Level1' not in col_map: col_map['Level1'] = df.columns[1]
            if 'Level2' not in col_map: col_map['Level2'] = df.columns[2]
            if 'Level3' not in col_map: col_map['Level3'] = df.columns[3]
            if 'Process' not in col_map: col_map['Process'] = df.columns[6]
            if 'Description' not in col_map: col_map['Description'] = df.columns[7]
        else:
            print("列数不足，无法继续")
            logger.error("列数不足，无法继续")
            return

    print(f"列映射: {col_map}")
    logger.info(f"列映射: {col_map}")

    # 重命名列
    df_renamed = df.rename(columns={
        col_map['CustomerReq']: 'CustomerReq',
        col_map['Level1']: 'Level1',
        col_map['Level2']: 'Level2',
        col_map['Level3']: 'Level3',
        col_map['Process']: 'Process',
        col_map['Description']: 'Description'
    })
    
    # 提取需要的列
    df = df_renamed[['CustomerReq', 'Level1', 'Level2', 'Level3', 'Process', 'Description']].copy()
    
    # 清理"功能过程"列中的无效关键字（这些应该在子过程描述中，而不是功能过程列）
    invalid_keywords = ['呈现', '查询', '保存', '输入', '校验', '输出']
    df.loc[df['Process'].isin(invalid_keywords), 'Process'] = None
    
    # 向下填充模块列和功能过程列（处理合并单元格）
    # 【关键】加上CustomerReq列，确保不同客户需求下的相同模块不会被合并
    df[['CustomerReq', 'Level1', 'Level2', 'Level3', 'Process']] = df[['CustomerReq', 'Level1', 'Level2', 'Level3', 'Process']].ffill()
    
    # 过滤掉可能是表头重复的行（例如值为"一级模块"的行）
    df = df[df['Level1'].astype(str).str.contains('一级模块', na=False) == False]
    
    # 关键：对数据进行排序，确保 Word 文档的顺序与逻辑结构一致
    # 这也确保了如果 Excel 乱序，生成的文档是规整的，且验证脚本也能通过（如果验证脚本也排序）
    # 注意：这里假设模块名称是字符串，排序可能按字典序。如果需要按原文件出现顺序排序但又要分组，
    # 则需要更复杂的逻辑（例如记录首次出现的索引）。
    # 鉴于通常需求是按 Excel 原序（如果原序就是分组好的），或者按模块聚类。
    # 为了稳妥，我们保持 Excel 原序的"分组"。
    # 即：不显式 sort_values，而是依赖 groupby(sort=False)。
    # 但是，如果 Excel 是乱序的 (A, B, A)，groupby(sort=False) 会输出 A, B (A被合并到第一个A)。
    # 这会导致 Word 内容顺序 = A(all), B(all).
    # 而 Excel 原序 = A1, B1, A2.
    # 验证脚本按原序读取: A1, B1, A2.
    # Word 文档: A1, A2, B1.
    # 于是验证失败。
    
    # 解决方案：强制验证脚本也按"分组后"的顺序来预期。
    # 或者，我们假设 Excel 本身就是排好序的。
    # 既然验证失败，说明 Excel 可能不是严格排序的，或者 groupby 改变了顺序。
    # 让我们在 converter 中不做改变（保持 groupby 聚合），但在 verify 中模拟这种聚合。
    
    # 创建Word文档
    doc = Document()
    
    # 状态变量，用于控制标题输出
    current_l1 = None
    current_l2 = None
    
    # 序号计数器
    idx_l1 = 0
    idx_l2 = 0
    idx_l3 = 0
    
    # 按模块分组 (CustomerReq, L1, L2, L3)
    # 【关键】加上CustomerReq确保不同客户需求下的相同模块不会被合并
    # 使用 groupby(sort=False) 保持 Excel 中的顺序
    module_groups = df.groupby(['CustomerReq', 'Level1', 'Level2', 'Level3'], sort=False)
    
    for (customer_req, l1, l2, l3), group_df in module_groups:
        # 1. 处理一级模块 (标题 3)
        if l1 != current_l1:
            idx_l1 += 1
            idx_l2 = 0 # 重置二级计数
            idx_l3 = 0 # 重置三级计数
            
            title_text = f"{l1}"
            add_styled_heading(doc, title_text, level=3)
            
            current_l1 = l1
            current_l2 = None # 重置二级模块状态
            
        # 2. 处理二级模块 (标题 4)
        if l2 != current_l2:
            idx_l2 += 1
            idx_l3 = 0 # 重置三级计数
            
            title_text = f"{l2}"
            add_styled_heading(doc, title_text, level=4)
            
            current_l2 = l2
            
        # 3. 处理三级模块 (标题 5)
        idx_l3 += 1
        title_text = f"{l3}"
        add_styled_heading(doc, title_text, level=5)
        
        # 4. 关键时序图/业务逻辑图 (标题 6)
        # 序号: L1.L2.L3.1
        title_text = "关键时序图/业务逻辑图"
        add_styled_heading(doc, title_text, level=6)
        
        para = doc.add_paragraph('无。')
        for run in para.runs:
            set_font(run)
            
        # 5. 功能描述 (标题 6)
        # 序号: L1.L2.L3.2
        title_text = "功能描述"
        add_styled_heading(doc, title_text, level=6)
        
        # 6. 整体功能列表
        # 获取该模块下所有唯一的功能过程
        processes = group_df['Process'].dropna().unique()
        # 过滤掉单纯的关键字（如果有）
        valid_processes = [p for p in processes if str(p).strip() not in ['呈现', '查询', '保存', '输入', '校验', '输出']]
        
        if valid_processes:
            summary_text = "　整体功能列表包含如下：" + "、".join(valid_processes) + "。"
            para = doc.add_paragraph(summary_text)
            for run in para.runs:
                set_font(run)
        
        # 7. 详细功能列表
        # 在当前三级模块组内，按功能过程分组
        process_groups = group_df.groupby('Process', sort=False)
        
        p_idx = 1
        for p_name, p_rows in process_groups:
            p_name_str = str(p_name).strip()
            if p_name_str in ['呈现', '查询', '保存', '输入', '校验', '输出']:
                continue
                
            # 输出功能过程标题 (正文格式，带序号)
            # 例如: 1.传输-传输管线系统链路数据呈现
            title_para = doc.add_paragraph(f"{p_idx}.{p_name_str}")
            for run in title_para.runs:
                set_font(run)
            p_idx += 1
            
            # 输出子过程描述
            for _, row in p_rows.iterrows():
                desc = row['Description']
                if pd.isna(desc):
                    continue
                
                # 拆分描述（如果一行包含多个步骤）
                lines = split_subprocess_description(desc)
                for line in lines:
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        set_font(run)

    # 确定输出路径
    if word_path is None:
        excel_file = Path(excel_path)
        word_path = excel_file.parent / f"{excel_file.stem}.docx"
    
    # 如果输出文件已存在，先删除
    if Path(word_path).exists():
        try:
            Path(word_path).unlink()
        except PermissionError:
            print(f"无法删除文件 {word_path}，请确保文件未被打开。")
            logger.error(f"无法删除文件 {word_path}，请确保文件未被打开。")
            return
    
    # 保存Word文档
    try:
        doc.save(word_path)
        print("Word文档已生成~")
        logger.info("Word文档已生成~")

        # 调用验证
        if perform_verify and verify_consistency:
            print("正在进行内容校对...")
            logger.info("正在进行内容校对...")
            verify_consistency(excel_path, word_path)

        # 打开文件
        if open_output:
            try:
                if hasattr(os, 'startfile'):
                    os.startfile(word_path)
                    print(f"已打开文件: {word_path}")
                    logger.info(f"已打开文件: {word_path}")
                else:
                    print("当前系统不支持自动打开文件")
                    logger.info("当前系统不支持自动打开文件")
            except Exception as e:
                print(f"无法自动打开文件: {e}")
                logger.error(f"无法自动打开文件: {e}")

    except Exception as e:
        print(f"保存Word文档失败: {e}")
        logger.exception(f"保存Word文档失败: {e}")


if __name__ == "__main__":
    logger.info("此脚本仅供 Web 服务内部调用，不支持直接命令行运行")
    logger.info("请通过 Streamlit 应用界面使用转换功能")
